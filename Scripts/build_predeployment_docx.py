"""Generate the pre-deployment checklist DOCX.

Run: python scripts/build_predeployment_docx.py
Outputs:
  - docs/PRE-DEPLOYMENT-CHECKLIST-2026-05-28.docx
  - %USERPROFILE%/Desktop/ROSRA-V2-PRE-DEPLOYMENT-CHECKLIST-2026-05-28.docx
"""

import os
import shutil
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_REPO = os.path.join("docs", "PRE-DEPLOYMENT-CHECKLIST-2026-05-28.docx")
OUT_DESKTOP = os.path.join(
    os.path.expanduser("~"), "Desktop",
    "ROSRA-V2-PRE-DEPLOYMENT-CHECKLIST-2026-05-28.docx"
)

UNHAB_BLUE = RGBColor(0x00, 0x6E, 0xB8)
DARK = RGBColor(0x1F, 0x2A, 0x44)
MUTED = RGBColor(0x55, 0x5E, 0x6C)


def style_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = UNHAB_BLUE


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = UNHAB_BLUE


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10.5)
    return table


def build():
    doc = Document()
    style_normal(doc)
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    add_title(doc, "ROSRA V2 — Pre-Deployment Checklist")
    add_subtitle(doc, f"Prepared: {date.today().isoformat()} • Owner: Boniface Michuki (Dev) • Programme: UN-Habitat ROSRA")

    # Status summary
    add_h2(doc, "1. Status summary")
    add_para(
        doc,
        "Application code is feature-complete and the latest hardening pass has been "
        "verified locally (build clean, login rate limiter active, antiforgery tokens "
        "in place on state-changing endpoints, Identity password-reset chain wired to "
        "the email service, permission cache invalidation on role changes). "
        "Remaining items before go-live are external dependencies (email provisioning, "
        "production Azure resources, cyber-sec sign-off) and a small number of open "
        "product decisions."
    )

    # Pending items table
    add_h2(doc, "2. Outstanding items before go-live")

    add_table(
        doc,
        ["#", "Item", "Owner / Responsible", "Status", "Action needed"],
        [
            ["1", "SMTP / email provisioning (production credentials)",
             "Ndeti — Boniface to receive credentials",
             "Blocked on input",
             "Ndeti to provide production SMTP host, port, sender mailbox, and app password / OAuth secret. Code path is wired via EmailService + IdentityEmailSender; only credentials are missing."],
            ["2", "Independent cybersecurity assessment",
             "UN-Habitat cyber-sec team (via Ndeti)",
             "Pending review",
             "Application to be handed to the UN-Habitat cyber-sec team via Ndeti for independent review. Sign-off required before production cut-over."],
            ["3", "Production deployment resources (Azure subscription, App Service plan, SQL tier, custom domain)",
             "Ndeti — Boniface to deploy",
             "Blocked on provisioning",
             "Ndeti to confirm production Azure resource group, App Service plan (size/region), Azure SQL tier, custom domain & TLS certificate, and required app settings (ADMIN_SEED_PASSWORD, AllowedHosts, SMTP settings)."],
            ["4", "Report generation flow — registration requirement",
             "Product Manager",
             "Open decision",
             "Current behaviour: users must register and log in before creating or downloading a ROSRA report. Open question to PM: keep registration mandatory, or allow anonymous report generation (e.g. for partners running quick scenarios)? Any change here also affects audit-trail and email notifications."],
            ["5", "Azure HTTPS-only, min TLS 1.2, production AllowedHosts",
             "Ndeti (ops) + Dev",
             "Pending Azure config",
             "Verify App Service is set to HTTPS Only = On, Minimum TLS = 1.2, and set AllowedHosts to the final production host via app settings. Base appsettings.json still has '*'."],
        ],
    )

    # Responsibilities matrix
    add_h2(doc, "3. Responsibility matrix")
    add_table(
        doc,
        ["Area", "Responsible", "Accountable", "Consulted"],
        [
            ["Email / SMTP provisioning", "Ndeti", "Boniface (Dev)", "Product Manager"],
            ["Cyber-sec assessment", "UN-Habitat cyber-sec team (via Ndeti)", "Boniface (Dev)", "Programme lead"],
            ["Azure production resources", "Ndeti (ops)", "Boniface (Dev)", "Programme lead"],
            ["Report generation flow & registration policy", "Product Manager", "Programme lead", "Dev (Boniface), UX"],
        ],
    )

    # Decision asks
    add_h2(doc, "4. Decisions needed (please confirm)")
    add_para(doc, "From Ndeti:", bold=True)
    for item in [
        "Production SMTP credentials and sender mailbox.",
        "Azure subscription / resource group / App Service plan size for production.",
        "Custom domain and TLS certificate.",
        "Date for cyber-sec team's independent review.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "From Product Manager:", bold=True)
    for item in [
        "Must end users register before generating a ROSRA report, or do we allow anonymous report generation? (Current behaviour: registration required.)",
    ]:
        add_bullet(doc, item)

    # Footer note
    add_h2(doc, "5. References")
    add_bullet(doc, "Current dev environment: https://rosra-dev.azurewebsites.net.")

    os.makedirs(os.path.dirname(OUT_REPO), exist_ok=True)
    doc.save(OUT_REPO)
    shutil.copyfile(OUT_REPO, OUT_DESKTOP)
    print(f"Saved: {OUT_REPO}")
    print(f"Saved: {OUT_DESKTOP}")


if __name__ == "__main__":
    build()
