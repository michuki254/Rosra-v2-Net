#!/usr/bin/env python3
"""Build the ROSRA V2 implementation handover DOCX.

The report is intentionally generated from code so the handover can be
regenerated after replacing screenshots, updating production metadata, or
adjusting the permission matrix.
"""

from __future__ import annotations

import sys
import zipfile
from datetime import date
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, "/tmp/rosra_docx_py")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except ImportError:  # pragma: no cover - only used by local builder execution
    Image = None


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "ROSRA_V2_Implementation_Workflow_Access_Data_Management_Handover_Report.docx"
PREPARED_BY = "Boniface Michuki"
REFERENCE_DOCX = Path(
    "/Users/michuki/Downloads/output pay/"
    "EARLDraft ROSRA_Methodology_and_Workflow_Overview_Annex4_5_6_WithScreenshots.docx"
)

BLUE = RGBColor(0x00, 0x9E, 0xDB)
NAVY = RGBColor(0x0B, 0x25, 0x45)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
GRAY = RGBColor(0x55, 0x65, 0x72)
LIGHT_BLUE_HEX = "EAF6F8"
LIGHT_GRAY_HEX = "F2F2F2"
PALE_ORANGE_HEX = "FFF4DE"
PALE_GREEN_HEX = "F7FBFC"
PALE_RED_HEX = "FFF4DE"
TABLE_HEADER_HEX = "0B6F83"
TABLE_BAND_HEX = "F7F7F7"
BORDER_HEX = "B8C2CC"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, after=1):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color or RGBColor(0, 0, 0), bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if widths:
        set_table_geometry(table, widths)

    set_row_cant_split(table.rows[0])
    set_row_repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_HEX)
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=max(font_size, 8.5), after=0)

    for row_idx, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row_data):
            if row_idx % 2 == 0:
                set_cell_shading(cells[i], TABLE_BAND_HEX)
            set_cell_text(cells[i], value, size=font_size, after=1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc, title, body, fill=LIGHT_BLUE_HEX, accent=BLUE):
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    if body:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.line_spacing = 1.12
        r2 = p2.add_run(body)
        set_run_font(r2, size=10, color=RGBColor(0x1F, 0x29, 0x37))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_para(doc, text, style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=11, color=color or RGBColor(0, 0, 0), bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1 and text != "Document Control":
        p.paragraph_format.page_break_before = True
    for run in p.runs:
        if level == 1:
            set_run_font(run, size=16, color=BLUE, bold=True)
        elif level == 2:
            set_run_font(run, size=13, color=BLUE, bold=True)
        else:
            set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=GRAY, italic=True)


def add_figure(doc, image_path, caption, max_width=6.3, max_height=4.6):
    path = Path(image_path)
    if not path.exists():
        add_callout(
            doc,
            "Screenshot slot",
            f"Insert screenshot: {caption}",
            fill=PALE_ORANGE_HEX,
            accent=ORANGE,
        )
        return

    width = max_width
    if Image is not None:
        with Image.open(path) as img:
            w_px, h_px = img.size
        ratio = w_px / h_px if h_px else 1
        width = min(max_width, max_height * ratio)
        width = max(width, 2.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        st.paragraph_format.space_after = Pt(6)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_p.add_run("ROSRA V2 Implementation Handover")
    set_run_font(r, size=9, color=GRAY)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run("Page ")
    set_run_font(r, size=9, color=GRAY)
    add_page_number(footer_p)

    return doc


def cover_page(doc):
    logo = ROOT / "docs/assets/un-habitat-cover-logo.png"
    cover_image = ROOT / "docs/assets/rosra-cover-home.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(str(logo), width=Inches(1.25))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("ROSRA V2")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Implementation, Workflow, Access, and Data Management Handover Report")
    set_run_font(r, size=20, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(
        "A plain-language implementation handover for the ROSRA team, covering what was "
        "implemented, how the system is hosted, how users and permissions work, "
        "how data is managed, and what the team should maintain."
    )
    set_run_font(r, size=12, color=GRAY)

    if cover_image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(str(cover_image), width=Inches(6.5))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run("ROSRA public landing page and workflow entry point.")
        set_run_font(r, size=8.5, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        f"Prepared by {PREPARED_BY} | For ROSRA / UN-Habitat project team | "
        f"Version 1.0 | {date.today():%Y-%m-%d}"
    )
    set_run_font(r, size=10, color=GRAY, bold=True)

    doc.add_section(WD_SECTION.NEW_PAGE)


def front_matter(doc):
    add_heading(doc, "Document Control", 1)
    add_table(
        doc,
        ["Document Field", "Value"],
        [
            ("Project", "Rapid Own-Source Revenue Analysis (ROSRA) V2"),
            ("Document type", "Implementation handover and operational overview"),
            ("Prepared by", PREPARED_BY),
            ("Prepared from", "Implementation and operational handover review"),
            ("Prepared for", "ROSRA / UN-Habitat project, operations, and implementation team"),
            ("Version", "1.0"),
            ("Date", date.today().strftime("%Y-%m-%d")),
            ("Status", "Comprehensive handover report for team review and use"),
        ],
        widths=[1.9, 4.6],
        font_size=9.5,
    )
    add_table(
        doc,
        ["Version", "Date", "Prepared / Updated By", "Purpose"],
        [
            ("1.0", date.today().strftime("%Y-%m-%d"), PREPARED_BY, "Initial comprehensive implementation handover."),
        ],
        widths=[0.8, 1.0, 2.0, 2.7],
    )
    add_table(
        doc,
        ["Reviewed source", "Use in this document"],
        [
            ("README.md", "Project purpose, capabilities, technology stack, and operating notes."),
            ("docs/ARCHITECTURE.md", "Runtime shape, controllers, services, data model, authorization, localization, and security controls."),
            ("docs/SETUP_AND_DEPLOYMENT.md", "Local setup, database initialization, Azure hosting notes, configuration, and deployment checks."),
            ("docs/DATABASE_SCHEMA_AND_BACKUP.md", "Database table summary, JSON report fields, backup and restore guidance."),
            ("Future maintenance guide", "Maintenance rules, change recipes, verification checklist, and handover notes."),
            ("azure-pipelines.yml", "Azure DevOps build and deployment target for the hosted application."),
            ("Program.cs and Data/DbInitializer.cs", "Startup sequence, registered policies, roles, permissions, services, and seed behavior."),
        ],
        widths=[2.25, 4.25],
        font_size=9,
    )
    add_callout(
        doc,
        "Document intent",
        "This report is written for non-technical and mixed technical audiences. "
        "It explains the implemented system in operational language, while retaining "
        "enough implementation detail for future maintenance, audits, onboarding, and change control.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "ROSRA V2 is a server-rendered ASP.NET Core MVC application that supports "
        "subnational own-source revenue assessment. It allows authenticated users to "
        "create ROSRA reports, estimate revenue potential, complete stream-specific "
        "diagnostics, prioritize administrative gaps, select reform recommendations, "
        "submit reports for review, and export validated analysis outputs."
    )
    add_para(
        doc,
        "The implementation combines a structured report-authoring workflow with an "
        "administrative back office. The administrative side manages users, roles, "
        "permissions, data uploads, deleted reports, audit logs, email settings, "
        "solution cards, and system settings. The review workflow preserves snapshots "
        "and report artifacts so reviewers can work from stable versions of submitted analyses."
    )
    add_para(
        doc,
        "The application is implemented on .NET 8, ASP.NET Core MVC, Razor views, "
        "ASP.NET Core Identity, Entity Framework Core, SQL Server, Bootstrap, jQuery, "
        "Chart.js, Playwright/Chromium PDF generation, ClosedXML Excel export, and "
        "MailKit SMTP email delivery."
    )
    add_callout(
        doc,
        "Handover summary",
        "The system is not only a calculator. It is an end-to-end operational platform: "
        "user access, report authoring, validation workflow, evidence snapshots, exports, "
        "admin data management, role-based security, localization, and deployment automation "
        "have all been implemented and should be maintained together.",
        fill=PALE_GREEN_HEX,
        accent=RGBColor(0x02, 0x7A, 0x48),
    )

    add_heading(doc, "Purpose and Audience", 1)
    add_para(
        doc,
        "This report documents how ROSRA V2 has been implemented, how it operates, "
        "and what the project team should maintain. It explains what exists, "
        "how it works, where the responsibilities sit, and what must be checked when "
        "future changes are made."
    )
    add_bullets(
        doc,
        [
            "Project and programme teams can use it to understand the user journey and operational workflow.",
            "Administrators can use it to understand user access, role management, data uploads, audit logs, and backups.",
            "Reviewers can use it to understand submission, review, validation, snapshots, and artifacts.",
            "Future technical maintainers can use it to locate the main code areas and avoid changing one part of the workflow without checking the related views, services, exports, and data fields.",
            "Security and audit reviewers can use it as a first map of implemented controls, data handling, and maintenance responsibilities.",
        ],
    )

    add_heading(doc, "Contents", 1)
    contents = [
        "Part I - System Overview and Implementation Summary",
        "Part II - User Access, Roles, and Permissions",
        "Part III - Application Workflow",
        "Part IV - Data and Database Management",
        "Part V - Technical Implementation and Maintenance Notes",
        "Annex 1 - Role and Permission Matrix",
        "Annex 2 - Database Table Summary",
        "Annex 3 - Hosting and Configuration Checklist",
        "Annex 4 - Screenshot Checklist",
        "Annex 5 - Backup and Restore Checklist",
        "Annex 6 - Maintenance and Handover Checklist",
    ]
    add_numbered(doc, contents)

    add_heading(doc, "Glossary", 1)
    add_table(
        doc,
        ["Term", "Plain-language meaning"],
        [
            ("ROSRA", "Rapid Own-Source Revenue Analysis, the revenue diagnostic tool implemented in the application."),
            ("OSR", "Own-source revenue collected directly by a local or subnational government."),
            ("Top-down estimate", "A benchmark-based quick estimate of current OSR performance and room to improve."),
            ("Bottom-up diagnostic", "A stream-by-stream analysis of administrative revenue gaps."),
            ("Report artifact", "A generated PDF or Excel file tracked by the system."),
            ("Snapshot", "A stable saved copy of a report at submission, validation, or pre-edit backup time."),
            ("Role", "A user group such as Admin, User, or Reviewer."),
            ("Permission", "A named capability that can be attached to a role, such as ExportReports or ReviewReports."),
            ("Soft delete", "A delete action that hides a record but keeps it recoverable until retention rules purge it."),
        ],
        widths=[1.7, 4.8],
        font_size=9.2,
    )


def part_i(doc, images):
    add_heading(doc, "Part I - System Overview and Implementation Summary", 1)
    add_para(
        doc,
        "ROSRA V2 is an implementation of a digital revenue diagnostic and decision-support "
        "workflow. The application guides local government and project users from a general "
        "own-source revenue benchmark into a more detailed administrative gap analysis, then "
        "into prioritisation, reform recommendations, report submission, review, validation, "
        "and exports."
    )
    add_figure(
        doc,
        images["workflow"],
        "Figure 1. ROSRA diagnostic workflow showing quick potential estimate, detailed diagnostic, prioritisation, reform guidance, and diagnostic report output.",
        max_height=3.6,
    )

    add_heading(doc, "1.1 What Was Implemented", 2)
    add_para(
        doc,
        "The system covers the public site, user "
        "authentication, report authoring interface, calculations, review workflow, exports, "
        "administration, reference data management, and deployment pipeline. The system is "
        "structured so the team can operate ROSRA reports rather than manually managing files "
        "outside the application."
    )
    add_table(
        doc,
        ["Implemented area", "What it does for the team"],
        [
            ("Public and account pages", "Supports entry to the platform, login, registration, profile, and language switching."),
            ("Report builder", "Allows users to create, save, edit, reload, and export ROSRA assessments."),
            ("Quick Potential Estimate", "Provides a benchmark-based early estimate using local profile inputs and WoFi/reference data."),
            ("Detailed Diagnostic", "Captures stream-level inputs for property tax, business licenses, user charges, mixed charges, and custom streams."),
            ("Prioritisation", "Helps users order streams and gaps based on scale, feasibility, and reform priority."),
            ("Recommendations", "Connects diagnostic results to solution cards and implementation guidance."),
            ("Review and validation", "Manages submission, review, revision, validation, snapshots, notes, and artifacts."),
            ("Administration", "Manages users, roles, permissions, data uploads, reports, deleted reports, solution cards, settings, and audit logs."),
            ("Exports", "Generates PDF and Excel outputs for analysis, presentation, and offline sharing."),
        ],
        widths=[2.15, 4.35],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Key Implementation Note",
        "The implemented workflow should be treated as one connected operating system. "
        "Changes to report fields, formulas, permissions, or status rules should be checked "
        "against save/load behavior, review snapshots, admin views, PDF exports, Excel exports, "
        "and localization.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )

    add_heading(doc, "1.2 Technology Stack", 2)
    add_table(
        doc,
        ["Layer", "Technology used", "Plain-language role"],
        [
            ("Application framework", ".NET 8 / ASP.NET Core MVC", "Runs the web application and server-side pages."),
            ("Views", "Razor views and partial views", "Build the multi-step report interface and admin pages."),
            ("Authentication", "ASP.NET Core Identity", "Stores users, roles, password hashes, login state, and account data."),
            ("Authorization", "Roles plus custom permission policies", "Controls who can access reports, admin tools, and review functions."),
            ("Database", "SQL Server with Entity Framework Core", "Stores users, reports, permissions, workflow records, reference data, and audit logs."),
            ("Frontend", "Bootstrap, jQuery, Chart.js, Select2, Font Awesome", "Provides responsive layout, forms, charts, dropdowns, and icons."),
            ("PDF export", "Playwright/Chromium and QuestPDF", "Renders report outputs and PDF-ready views."),
            ("Excel export", "ClosedXML", "Creates spreadsheet exports."),
            ("Email", "MailKit / SMTP", "Sends account and workflow notifications when configured."),
            ("Localization", "ASP.NET Core localization resources", "Supports English, French, and Spanish labels."),
        ],
        widths=[1.45, 2.0, 3.05],
        font_size=8.4,
    )

    add_heading(doc, "1.3 Hosting and Deployment", 2)
    add_para(
        doc,
        "The repository includes an Azure DevOps pipeline that builds the application and "
        "deploys it to Azure App Service. The pipeline target is the App Service named "
        "rosra-dev in resource group UNHAB-PRD-ROSRA, using the Azure service connection "
        "Azure-RosraDev. If no custom domain is configured, the default Azure App Service "
        "host would normally be based on the app name."
    )
    add_table(
        doc,
        ["Hosting item", "Implemented value / note"],
        [
            ("Pipeline file", "azure-pipelines.yml"),
            ("Build trigger", "Pushes to main, excluding documentation-only paths."),
            ("Build agent", "windows-latest"),
            ("Deployment method", "Azure Web App zip deploy."),
            ("App Service name", "rosra-dev"),
            ("Resource group", "UNHAB-PRD-ROSRA"),
            ("Subscription label in pipeline comments", "UN-HABITAT-PROD"),
            ("Minimum runtime setting", "ASPNETCORE_ENVIRONMENT=Production"),
            ("Database setting", "CONNECTION_STRING or Azure App Service connection string."),
            ("Persistent storage requirement", "Data Protection keys and Playwright browser cache must survive restarts."),
        ],
        widths=[2.05, 4.45],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Hosting control",
        "The exact live URL, custom domain, Azure SQL database name, and production owner "
        "should be recorded in the team's operational register. The source code identifies "
        "the App Service target and deployment method, but the final public domain may be "
        "configured outside the repository.",
        fill=PALE_ORANGE_HEX,
        accent=ORANGE,
    )
    add_figure(
        doc,
        images["login"],
        "Figure 2. User-facing entry point: account login and access into the ROSRA tool.",
        max_height=4.1,
    )


def part_ii(doc, images):
    add_heading(doc, "Part II - User Access, Roles, and Permissions", 1)
    add_para(
        doc,
        "ROSRA V2 uses ASP.NET Core Identity for user accounts. The implemented access "
        "model combines normal authenticated-user checks, broad role checks, and named "
        "permission policies. This gives the team a practical way to separate normal users, "
        "reviewers, and administrators."
    )
    add_figure(
        doc,
        images["register"],
        "Figure 3. Account creation and onboarding screen used by new ROSRA users.",
        max_height=4.4,
    )

    add_heading(doc, "2.1 Account Workflow", 2)
    add_numbered(
        doc,
        [
            "A user opens the public ROSRA site and signs in or creates an account.",
            "Local email/password users are stored through ASP.NET Core Identity.",
            "Optional Microsoft Entra ID sign-in is enabled only when tenant and client settings are configured.",
            "New registered users are assigned the User role by default.",
            "Administrators can manage user roles, account activation, and access through the admin area.",
            "The initial admin account admin@rosra.com is created only when ADMIN_SEED_PASSWORD is configured.",
        ],
    )

    add_heading(doc, "2.2 Implemented Roles", 2)
    add_table(
        doc,
        ["Role", "Primary purpose", "Typical responsibilities"],
        [
            ("User", "Report author", "Create reports, save/edit own reports, submit reports, export outputs, use dashboard."),
            ("Reviewer", "Assessment reviewer", "Review submitted reports, add notes, validate or request revision, access snapshots and artifacts."),
            ("Admin", "System administrator", "Manage users, roles, permissions, all reports, uploads, solution cards, email settings, system settings, backups, and audit logs."),
        ],
        widths=[1.1, 1.7, 3.7],
        font_size=8.8,
    )

    add_heading(doc, "2.3 Access Control Implementation", 2)
    add_para(
        doc,
        "The application has three access-control patterns. First, ordinary authenticated "
        "pages use an authenticated-user requirement. Second, sensitive administrative and "
        "review pages use role checks, such as Admin or Admin/Reviewer. Third, permission "
        "policies are registered for named capabilities such as UploadPeerSNGData, "
        "ExportReports, ReviewReports, and ManagePermissions."
    )
    add_callout(
        doc,
        "Access Control Note",
        "The Admin area is protected by the Admin role at controller level. The Submission "
        "workflow is protected by authenticated access plus Admin/Reviewer role checks on "
        "review actions. The custom permission system is seeded and available for fine-grained "
        "checks; future maintainers should confirm that any newly introduced permission name "
        "is both seeded and registered as an authorization policy.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )
    add_table(
        doc,
        ["Category", "Examples of permissions", "Default role intent"],
        [
            ("Reports", "ViewReports, CreateReports, EditReports, DeleteReports, ExportReports", "User, Reviewer, Admin"),
            ("All-report administration", "EditAllReports, DeleteAllReports", "Admin"),
            ("User management", "ViewUsers, CreateUsers, EditUsers, DeleteUsers, ManageUserRoles, ActivateDeactivateUsers", "Admin"),
            ("Data management", "UploadPeerSNGData, UploadCountryData, ViewDataLibrary, DeleteUploadedData", "Admin or assigned data manager"),
            ("Role management", "ViewRoles, CreateRoles, EditRoles, DeleteRoles, ManagePermissions", "Admin"),
            ("Dashboards", "ViewDashboard, ViewAdminDashboard, ViewAnalytics", "User for own dashboard; Admin for admin dashboard and analytics"),
            ("Review workflow", "SubmitReports, ReviewReports, ValidateReports, AssignReviewers, ViewReviewNotes, AddReviewNotes", "User submits; Reviewer and Admin review/validate"),
            ("Artifacts and snapshots", "ViewAnalysisSnapshots, AccessReportArtifacts", "Reviewer and Admin"),
            ("Advanced review operations", "BulkValidate, ReRunCalculations, UnlockValidatedReports", "Reviewer/Admin for selected operations; Admin for unlock"),
        ],
        widths=[1.45, 3.0, 2.05],
        font_size=8.0,
    )

    add_heading(doc, "2.4 Administrative Responsibilities", 2)
    add_para(
        doc,
        "Administrators are responsible for controlling who can use the system and what they "
        "can do. The admin area also gives operational visibility into reports, deleted "
        "reports, audit logs, data uploads, analytics, system settings, email settings, and "
        "solution-card content."
    )
    add_bullets(
        doc,
        [
            "Create, activate, deactivate, and update user accounts.",
            "Assign or remove roles, including Reviewer and Admin.",
            "Maintain permission assignments in the role/permission screen.",
            "Review report inventory and deleted reports.",
            "Manage reference uploads and confirm upload history.",
            "Review audit logs for sensitive actions.",
            "Maintain email and system settings without committing secrets into source control.",
        ],
    )
    add_figure(
        doc,
        images["admin_data"],
        "Figure 4. Admin area navigation showing reports, data management, solution cards, users/access, audit log, and configuration sections.",
        max_height=4.25,
    )
    add_callout(
        doc,
        "Security Action Required",
        "Any deployment token, password, SMTP credential, connection string, or personal "
        "access token must be stored in approved secret storage, not in shared documents or "
        "source-controlled files. If any token has been exposed during development, rotate it "
        "and update the Azure DevOps service connection or platform secret.",
        fill=PALE_RED_HEX,
        accent=RGBColor(0xB4, 0x23, 0x18),
    )


def part_iii(doc, images):
    add_heading(doc, "Part III - Application Workflow", 1)
    add_para(
        doc,
        "The user workflow is designed to move from entry and report creation into analysis, "
        "prioritisation, recommendations, submission, review, validation, and export. Each "
        "step stores data so users can save, return, review, and generate outputs."
    )

    add_heading(doc, "3.1 End-to-End User Journey", 2)
    add_table(
        doc,
        ["Stage", "User action", "System behavior"],
        [
            ("Access", "User signs in or registers.", "Identity authenticates the user and applies roles."),
            ("Start", "User starts a new analysis or opens a saved report.", "A RosraReport record is created or loaded."),
            ("Profile", "User enters country, government unit, financial year, currency, OSR, population, and economic profile.", "Profile data is used by top-down and downstream calculations."),
            ("Quick estimate", "User runs the benchmark-based estimate.", "The system calculates current OSR, estimated potential, and room to improve."),
            ("Detailed diagnostic", "User enters stream-specific data.", "The system decomposes gaps by stream and administrative area."),
            ("Prioritisation", "User ranks streams and gaps.", "Selected priorities are stored for recommendations and exports."),
            ("Recommendations", "User selects solution cards and implementation guidance.", "Solutions and action planning data are saved with the report."),
            ("Submission", "User submits for review.", "Status changes, snapshot is created, and reviewers can inspect a stable version."),
            ("Review/validation", "Reviewer adds notes, requests revision, or validates.", "Workflow status, notes, validation user, and snapshots are tracked."),
            ("Export", "User or reviewer exports PDF/Excel.", "Generated artifacts are tracked with metadata."),
        ],
        widths=[1.3, 2.4, 2.8],
        font_size=8.1,
    )

    add_callout(
        doc,
        "Workflow Rule",
        "The implemented report status path is: Draft -> Submitted -> UnderReview -> "
        "NeedsRevision -> Validated. The team should not manually bypass this sequence "
        "without understanding the snapshot, artifact, and review-note effects.",
        fill=PALE_ORANGE_HEX,
        accent=ORANGE,
    )

    add_heading(doc, "3.2 Creating and Starting a Report", 2)
    add_para(
        doc,
        "The report builder lives under the main ROSRA report flow. Users can start a new "
        "analysis, complete profile information, and move through the tool using tabbed or "
        "step-based navigation. Report data can be saved and reloaded from the dashboard."
    )
    add_figure(
        doc,
        images["quick_profile"],
        "Figure 5. Local government profile and Quick Potential Estimate input screen.",
        max_height=4.3,
    )

    add_heading(doc, "3.3 Quick Potential Estimate", 2)
    add_para(
        doc,
        "The Quick Potential Estimate is the early benchmark view. It gives the user an "
        "indicative assessment of current OSR performance and room to improve before the "
        "user enters detailed stream-specific data. The estimator uses embedded WoFi and "
        "reference datasets where available, and user-provided inputs where needed."
    )
    add_figure(
        doc,
        images["quick_results"],
        "Figure 6. Quick Potential Estimate output showing current OSR, estimated potential, and room to improve.",
        max_height=4.5,
    )

    add_heading(doc, "3.4 Detailed Diagnostic and Gap Analysis", 2)
    add_para(
        doc,
        "The Detailed Diagnostic captures revenue stream inputs and converts them into "
        "administrative gap estimates. Implemented streams include property tax, business "
        "licenses, short-term user charges, long-term user charges, mixed user charges, and "
        "generic/custom streams."
    )
    add_table(
        doc,
        ["Stream / area", "Purpose in the diagnostic"],
        [
            ("Property tax", "Estimate coverage, compliance, valuation, and related gaps for property-based revenue."),
            ("Business license", "Estimate business registration, compliance, and liability gaps."),
            ("Short-term user charges", "Model short-cycle usage, actual payment, and revenue leakage."),
            ("Long-term user charges", "Model monthly or long-cycle user charges and collection gaps."),
            ("Mixed user charges", "Model revenue streams with both short-term and longer-term users."),
            ("Generic streams", "Allow custom non-property revenue streams using the generic gap calculation service."),
            ("Total estimate", "Aggregate gap outputs across selected streams for report-level interpretation."),
        ],
        widths=[2.0, 4.5],
        font_size=8.6,
    )
    add_figure(
        doc,
        images["gap_analysis"],
        "Figure 7. Stream-specific diagnostic screen showing gap-analysis results and visual decomposition.",
        max_height=4.5,
    )

    add_heading(doc, "3.5 Prioritisation", 2)
    add_para(
        doc,
        "Prioritisation turns diagnostic outputs into a sequence for action. Users can assess "
        "which streams and gap types should be addressed first. This is important because a "
        "large revenue gap is not automatically the first reform priority; feasibility, "
        "administrative capacity, and policy context also matter."
    )
    add_figure(
        doc,
        images["prioritisation"],
        "Figure 8. Prioritisation view showing selected revenue streams and gap priority sequence.",
        max_height=4.2,
    )

    add_heading(doc, "3.6 Recommendations and Implementation Guidance", 2)
    add_para(
        doc,
        "Recommendations connect the diagnostic results to solution cards. Solution cards "
        "are stored and managed in the application so the team can maintain consistent reform "
        "guidance across reports. The recommendation workflow supports solution cards, "
        "timeline view, and implementation progress tracking."
    )
    add_figure(
        doc,
        images["recommendations"],
        "Figure 9. Recommendations screen showing solution cards and implementation guidance.",
        max_height=4.2,
    )

    add_heading(doc, "3.7 Review, Validation, and Exports", 2)
    add_para(
        doc,
        "When a report is submitted, the review workflow creates a stable snapshot and allows "
        "reviewers to add notes or validate the assessment. Exports can be generated as PDF "
        "or Excel files. Report artifacts are recorded in the database with metadata such as "
        "file name, file type, generation date, and generating user."
    )
    add_figure(
        doc,
        images["export_modal"],
        "Figure 10. Export modal for generating the ROSRA diagnostic report.",
        max_height=3.4,
    )
    add_table(
        doc,
        ["Export", "Implemented tool", "Use"],
        [
            ("PDF report", "Playwright/Chromium-backed HTML-to-PDF and report export services", "Formal diagnostic report output for sharing and review."),
            ("Excel workbook", "ClosedXML through ExcelExportService", "Structured data export for offline analysis."),
            ("Admin data export", "Admin data-management export", "Filtered report data for external analysis and monitoring."),
            ("Solution-card JSON", "Admin solution-card export", "Backup or migration of solution card content."),
        ],
        widths=[1.7, 2.4, 2.4],
        font_size=8.6,
    )


def part_iv(doc, images):
    add_heading(doc, "Part IV - Data and Database Management", 1)
    add_para(
        doc,
        "ROSRA V2 uses SQL Server through Entity Framework Core. The database stores user "
        "accounts, roles, permissions, reports, workflow records, uploaded data history, "
        "reference datasets, solution cards, email settings, audit logs, and generated "
        "artifact metadata. The database stores report data; generated files and platform "
        "secrets must also be managed outside SQL backups."
    )
    add_figure(
        doc,
        images["admin_data"],
        "Figure 11. Admin data-management page used to browse, filter, inspect, and export assessment records.",
        max_height=4.2,
    )

    add_heading(doc, "4.1 Main Data Areas", 2)
    add_table(
        doc,
        ["Data area", "Main tables", "What the team should understand"],
        [
            ("Users and access", "AspNetUsers, AspNetRoles, AspNetUserRoles, Permissions, RolePermissions", "Controls accounts, roles, and capabilities."),
            ("Reports", "RosraReports", "Central report record, including profile, workflow status, and JSON tab payloads."),
            ("Review workflow", "ReviewNotes, AnalysisSnapshots, ReportArtifacts", "Preserves submitted versions, reviewer notes, validation records, and generated files metadata."),
            ("Reference data", "Country, DB_Countries, DB_Frontiers, Peers_SNG, UserPeerSngs, DataUploadHistory", "Supports country lookup, benchmark data, peer analysis, and upload history."),
            ("Administration", "AuditLogs, EmailSettings, EmailLogs, SystemSettings", "Supports operational settings, email delivery, and accountability."),
            ("Solution library", "SolutionCards, SolutionCardHistory, CardSets", "Stores recommendation cards and change history."),
            ("Migrations", "__EFMigrationsHistory", "Tracks applied database schema migrations."),
        ],
        widths=[1.5, 2.35, 2.65],
        font_size=7.8,
    )

    add_heading(doc, "4.2 How Reports Are Stored", 2)
    add_para(
        doc,
        "The central report table is RosraReports. A single row represents a ROSRA assessment. "
        "Simple metadata such as title, country, city, currency, financial year, owner, "
        "status, and timestamps are stored in normal columns. Complex tab-level form data is "
        "stored as JSON strings so the report structure can evolve without requiring a new "
        "database table for every form section."
    )
    doc.add_page_break()
    add_table(
        doc,
        ["RosraReports field group", "Examples", "Purpose"],
        [
            ("Identity", "Id, PublicId, Title", "Internal record ID, URL-safe report ID, and report title."),
            ("Ownership", "UserId, LastModifiedByUserId", "Report author and last editor."),
            ("Location/profile", "Country, Region, City, Government unit, Currency, FinancialYear", "Defines the assessment area and reporting period."),
            ("Financial profile", "ActualOsr, BudgetedOsr, Population, GdpPerCapita", "Inputs used for estimates and report context."),
            ("Gap JSON", "PropertyTaxData, LicenseData, ShortTermUserChargeData, LongTermUserChargeData, MixedUserChargeData, GenericStreamsData", "Detailed tab-level diagnostic data."),
            ("Recommendations JSON", "RootCauses, ActionItems, PrioritizationData, SelectedSolutionsData, ImplementationProgressData", "Causes, priorities, solution choices, and implementation tracking."),
            ("Workflow", "Status, CompletionLevel, SubmissionVersion, SubmittedAt, ValidatedAt, ReviewerUserId", "Submission, review, revision, and validation state."),
            ("Lifecycle", "IsDeleted, DeletedAt, DeletedByUserId, IsArchived, RowVersion", "Soft-delete, archive, and concurrency control."),
        ],
        widths=[1.55, 2.35, 2.6],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Data Management Note",
        "When changing a form field or calculation, check all places that use the same data: "
        "Razor form, JavaScript state collector, C# save/load mapping, JSON field shape, "
        "view mode, review snapshot, PDF export, Excel export, and localization labels.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )

    add_heading(doc, "4.3 Formula and Calculation Locations", 2)
    add_table(
        doc,
        ["Area", "Main implementation location", "Reason it matters"],
        [
            ("Property tax and business license calculations", "Services/GapCalculator.cs and related Razor partials", "Server and live UI results must remain consistent."),
            ("Generic stream calculations", "Services/GapCalculationService.cs and Controllers/Api/GapsApiController.cs", "Supports custom stream calculation through a server API."),
            ("WoFi/top-down potential estimate", "Services/WofiPotentialEstimator.cs and embedded Data/Wofi JSON", "Drives benchmark-based potential estimates."),
            ("Live UI calculations", "Views/Shared/_GapAnalysis*.cshtml and JavaScript modules", "Users see immediate results while entering form data."),
            ("Recommendation behavior", "wwwroot/js/recommendationsModule.js and SolutionCards data", "Connects gaps and priorities to reform guidance."),
            ("Exports", "ReportExportService, ExcelExportService, HtmlToPdfService", "Outputs must reproduce the saved report data accurately."),
        ],
        widths=[1.85, 2.45, 2.2],
        font_size=8.0,
    )

    add_heading(doc, "4.4 Reference Data and Upload Management", 2)
    add_para(
        doc,
        "Reference datasets are seeded at startup and can also be managed through the admin "
        "area. Core seed data includes countries, peer SNG data, WoFi data, country/state "
        "administrative divisions, and solution cards. Upload history is retained so the team "
        "can see what data was loaded and when."
    )
    add_figure(
        doc,
        images["admin_priority"],
        "Figure 12. Admin data-management view with expanded prioritisation information.",
        max_height=4.2,
    )
    add_table(
        doc,
        ["Reference source", "Where it lives", "Operational note"],
        [
            ("Country data", "Data/SeedData/countrydata.json and CountryStatesData.cs", "Keep country and administrative-division updates coordinated."),
            ("Peer SNG data", "Data/SeedData/peersng.json and Peers_SNG table", "Used for peer comparison and frontier analysis."),
            ("WoFi reference data", "Data/Wofi/*.json embedded resources", "Travels with compiled application; update project file if adding new embedded files."),
            ("User-entered peer data", "UserPeerSngs table", "Separated from admin-managed reference data so user entries do not overwrite global data."),
            ("Solution cards", "wwwroot/js/solutionsData-*.js, SolutionCards, SolutionCardHistory", "Admin-maintained recommendation library with change history."),
            ("Upload history", "DataUploadHistory table", "Track reference-data uploads and support operational accountability."),
        ],
        widths=[1.75, 2.3, 2.45],
        font_size=8.0,
    )

    add_heading(doc, "4.5 Backup, Restore, and Retention", 2)
    add_para(
        doc,
        "The built-in Excel exports are useful for analysis but are not full database backups. "
        "Full recovery requires a SQL Server backup or Azure SQL bacpac export, plus separate "
        "care for files and secrets that are outside the database."
    )
    add_table(
        doc,
        ["Item", "Protected by SQL backup?", "Required handling"],
        [
            ("Database rows", "Yes", "Use SQL Server .bak or Azure SQL .bacpac for full backup and restore."),
            ("Generated report files", "No", "Back up artifact file storage separately from database metadata."),
            ("Data Protection keys", "No", "Back up persistent key directory so auth cookies and encrypted settings survive restore."),
            ("Environment variables and secrets", "No", "Store in Azure App Service settings, Azure DevOps secrets, or approved secret manager."),
            ("Uploaded source files", "Depends on storage design", "Confirm whether original files are retained outside SQL."),
            ("Playwright browser binaries", "No backup needed", "Can be reinstalled; confirm install completes before relying on PDF export."),
        ],
        widths=[1.8, 1.55, 3.15],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Retention control",
        "Soft-deleted reports are excluded from normal views and can be restored before purge. "
        "The retention service is designed to purge expired deleted reports and related child "
        "records after the configured retention window. Always create a proper database backup "
        "before manual retention cleanup in production.",
        fill=PALE_GREEN_HEX,
        accent=RGBColor(0x02, 0x7A, 0x48),
    )


def part_v(doc, images):
    add_heading(doc, "Part V - Technical Implementation and Maintenance Notes", 1)
    add_para(
        doc,
        "This section is written for the next technical maintainer, technical project manager, or system "
        "owner who needs to maintain ROSRA V2. It avoids unnecessary code detail but identifies "
        "where responsibilities sit in the repository and what must be checked after changes."
    )
    add_figure(
        doc,
        images["admin_data"],
        "Figure 13. Admin module used for ongoing operational maintenance and monitoring.",
        max_height=4.0,
    )

    add_heading(doc, "5.1 Repository Map", 2)
    add_table(
        doc,
        ["Path", "Plain-language responsibility"],
        [
            ("Controllers/", "Request handlers for public pages, reports, dashboard, admin, submission, account, culture, and API endpoints."),
            ("Views/", "Razor pages and partials that render the user interface."),
            ("Models/", "Domain objects, view models, enums, and data shapes."),
            ("Data/", "Database context, initializer, embedded seed data, and WoFi reference data."),
            ("Services/", "Business services for workflow, exports, snapshots, artifacts, email, validation, seeding, retention, and calculations."),
            ("Authorization/", "Custom permission requirement, handler, and attribute support."),
            ("Infrastructure/", "Framework helpers such as comma-tolerant numeric model binding."),
            ("Resources/", "Localization files for English, French, and Spanish."),
            ("Migrations/", "Entity Framework Core database schema migrations."),
            ("wwwroot/", "Static assets, JavaScript modules, CSS, images, fonts, and public docs."),
            ("docs/", "Architecture, deployment, database, security, and handover documentation."),
            ("tools/", "SQL and data-processing helper scripts."),
        ],
        widths=[1.6, 4.9],
        font_size=8.2,
    )

    add_heading(doc, "5.2 Main Controllers", 2)
    add_table(
        doc,
        ["Controller", "Implemented responsibility"],
        [
            ("HomeController", "Public pages, job aid, sample reports, and admin or technical guide pages."),
            ("AccountController", "Registration, login, logout, profile, password changes, external sign-in callback."),
            ("RosraController", "Main report editor, save/edit/view operations, estimator endpoints, exports, print views, peer SNG analysis."),
            ("DashboardController", "User dashboard, report view, archive, delete, restore, and bulk actions."),
            ("SubmissionController", "Submission queue, review, validation, rejection, notes, snapshots, artifacts, unlock workflow."),
            ("AdminController", "Admin dashboards, users, roles, permissions, reports, deleted reports, audit logs, uploads, backups, settings, solution cards, analytics."),
            ("CultureController", "Language switching."),
            ("GapsApiController", "Generic stream gap calculation API."),
            ("SolutionsApiController", "Solution-card API for active and grouped card data."),
        ],
        widths=[1.75, 4.75],
        font_size=8.1,
    )

    add_heading(doc, "5.3 Main Services", 2)
    add_table(
        doc,
        ["Service", "Implemented responsibility"],
        [
            ("SubmissionService", "Moves reports through submission and review states."),
            ("ValidationService", "Checks completion and workflow rules before validation."),
            ("SnapshotService", "Captures stable report states for submission, validation, and pre-edit backup."),
            ("ArtifactService", "Creates and tracks generated report artifacts."),
            ("HtmlToPdfService", "Uses a long-lived Playwright browser for PDF rendering."),
            ("ReportExportService", "Builds localized report export content."),
            ("ExcelExportService", "Builds Excel report and admin exports."),
            ("EmailService", "Sends SMTP notifications and logs delivery attempts."),
            ("DataRetentionService", "Purges expired soft-deleted report data after retention period."),
            ("WofiPotentialEstimator", "Loads embedded WoFi data and estimates top-down OSR potential."),
            ("GapCalculationService / GapCalculator", "Performs gap calculation support for stream diagnostics."),
            ("SolutionCardSeeder / SampleReportSeeder", "Seeds solution-card and sample report content."),
        ],
        widths=[2.15, 4.35],
        font_size=8.0,
    )

    add_heading(doc, "5.4 Startup and Configuration", 2)
    add_numbered(
        doc,
        [
            "Resolve the database connection string from environment variables, Azure connection strings, or appsettings fallback.",
            "Register Entity Framework Core with SQL Server and retry-on-failure.",
            "Register Identity, roles, Data Protection keys, localization, MVC, antiforgery, authorization, sessions, memory cache, and rate limiting.",
            "Optionally register Microsoft Entra ID login when tenant and client settings are present.",
            "Register report, export, snapshot, artifact, email, validation, retention, calculation, and estimator services.",
            "Install Playwright Chromium in the background for PDF generation.",
            "Initialize the database through DbInitializer, applying migrations and seeding reference data and roles.",
            "Configure static files, request localization, security headers, authentication, authorization, session, and MVC routing.",
        ],
    )

    add_heading(doc, "5.5 Security Controls Implemented", 2)
    add_table(
        doc,
        ["Control", "Implementation summary"],
        [
            ("Authentication", "ASP.NET Core Identity local login with optional Entra ID sign-in."),
            ("Roles and permissions", "Admin, User, Reviewer roles plus seeded permission records and registered policies."),
            ("Cookies", "HTTP-only, Secure, SameSite Strict auth and session cookies with limited lifetime."),
            ("Data Protection", "Keys persisted to stable storage so cookies and encrypted values survive restarts."),
            ("Antiforgery", "Header-based antiforgery token support for AJAX requests."),
            ("Rate limiting", "Fixed-window API and login rate limiting."),
            ("Security headers", "Content type, frame, referrer, permissions policy, CSP, and cache controls."),
            ("HTTP methods", "Request method allow-list for GET, HEAD, POST, and OPTIONS."),
            ("Audit logs", "Report changes, status transitions, admin actions, and retention purges are logged."),
            ("Soft delete", "Reports are hidden by query filter and retained before purge."),
        ],
        widths=[1.65, 4.85],
        font_size=8.0,
    )

    add_callout(
        doc,
        "Technical Handover Note",
        "Before changing any report workflow, confirm that save/load behavior, dashboard views, "
        "admin views, review snapshots, artifacts, PDF export, Excel export, localization, "
        "and permission checks still work. The application is integrated across several layers; "
        "a small form change can affect multiple outputs.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )

    add_heading(doc, "5.6 Minimum Verification After Changes", 2)
    add_table(
        doc,
        ["Change area", "Minimum verification"],
        [
            ("General code change", "Run dotnet build and confirm the app starts."),
            ("Report field change", "Create, save, reload, edit, view, print, export PDF, export Excel, and inspect old reports."),
            ("Formula change", "Compare live UI calculation with server/export calculation."),
            ("Permission change", "Verify Admin, User, Reviewer, and unauthorized access paths."),
            ("Database migration", "Apply to local/disposable database and inspect generated migration and model snapshot."),
            ("Seed data change", "Run startup seeding on local database and confirm idempotency."),
            ("Export change", "Verify browser print view, PDF output, Excel output, and chart/image inclusion."),
            ("Localization change", "Check English, French, and Spanish labels where visible."),
            ("Admin function change", "Confirm antiforgery protection, audit logging, and admin-only access."),
        ],
        widths=[1.8, 4.7],
        font_size=8.0,
    )


def annexes(doc):
    add_heading(doc, "Annex 1 - Role and Permission Matrix", 1)
    add_para(
        doc,
        "The table below summarizes the implemented permissions and their intended role "
        "assignment based on seeding logic and the application design. The Admin role receives "
        "all core permissions. The User role receives basic report permissions and SubmitReports. "
        "The Reviewer role receives report, review, validation, snapshot, artifact, bulk validate, "
        "and recalculation permissions."
    )
    add_para(
        doc,
        "Use this annex when onboarding administrators, reviewing account access, or checking "
        "whether a user should be able to create, review, validate, export, upload, or manage "
        "system data. The matrix should be checked against the live Admin roles and permissions "
        "screen before any external audit or production access review.",
    )
    add_callout(
        doc,
        "Permission maintenance note",
        "When a new permission is introduced, it should be added in three places: the "
        "authorization policy list, the database seeding logic, and the Admin role/permission "
        "assignment. The protected controller action or page should then be tested with Admin, "
        "Reviewer, User, and unauthenticated access.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )
    permission_rows = [
        ("Reports", "ViewReports", "View reports", "User, Reviewer, Admin"),
        ("Reports", "CreateReports", "Create new reports", "User, Reviewer, Admin"),
        ("Reports", "EditReports", "Edit own reports", "User, Reviewer, Admin"),
        ("Reports", "EditAllReports", "Edit any user's reports", "Admin"),
        ("Reports", "DeleteReports", "Delete own reports", "User, Reviewer, Admin"),
        ("Reports", "DeleteAllReports", "Delete any user's reports", "Admin"),
        ("Reports", "ExportReports", "Export PDF/Excel", "User, Reviewer, Admin"),
        ("User Management", "ViewUsers", "View user list", "Admin"),
        ("User Management", "CreateUsers", "Create users", "Admin"),
        ("User Management", "EditUsers", "Edit users", "Admin"),
        ("User Management", "DeleteUsers", "Delete users", "Admin"),
        ("User Management", "ManageUserRoles", "Assign/remove roles", "Admin"),
        ("User Management", "ActivateDeactivateUsers", "Activate/deactivate users", "Admin"),
        ("Data Management", "UploadPeerSNGData", "Upload peer SNG data", "Admin or assigned data manager"),
        ("Data Management", "UploadCountryData", "Upload country data", "Admin or assigned data manager"),
        ("Data Management", "ViewDataLibrary", "View data library", "Admin or assigned data manager"),
        ("Data Management", "DeleteUploadedData", "Delete uploaded data", "Admin"),
        ("Role Management", "ViewRoles", "View roles", "Admin"),
        ("Role Management", "CreateRoles", "Create roles", "Admin"),
        ("Role Management", "EditRoles", "Edit roles", "Admin"),
        ("Role Management", "DeleteRoles", "Delete roles", "Admin"),
        ("Role Management", "ManagePermissions", "Assign permissions", "Admin"),
        ("Dashboard", "ViewDashboard", "View user dashboard", "User, Reviewer, Admin"),
        ("Dashboard", "ViewAdminDashboard", "View admin dashboard", "Admin"),
        ("Dashboard", "ViewAnalytics", "View analytics", "Admin"),
        ("Assessment Review", "SubmitReports", "Submit report for review", "User, Reviewer, Admin"),
        ("Assessment Review", "ReviewReports", "Review submitted reports", "Reviewer, Admin"),
        ("Assessment Review", "ValidateReports", "Validate or approve reports", "Reviewer, Admin"),
        ("Assessment Review", "UnlockValidatedReports", "Unlock validated report for revision", "Admin"),
        ("Assessment Review", "AssignReviewers", "Assign reviewers", "Admin"),
        ("Assessment Review", "ViewReviewNotes", "View review notes", "Reviewer, Admin"),
        ("Assessment Review", "AddReviewNotes", "Add review notes", "Reviewer, Admin"),
        ("Assessment Review", "ViewAnalysisSnapshots", "View analysis snapshots", "Reviewer, Admin"),
        ("Assessment Review", "AccessReportArtifacts", "Access generated artifacts", "Reviewer, Admin"),
        ("Assessment Review", "BulkValidate", "Bulk validate reports", "Reviewer, Admin"),
        ("Assessment Review", "ReRunCalculations", "Re-run calculations", "Reviewer, Admin"),
        ("Solution Management", "ManageSolutionCards", "Manage solution cards", "Seeded permission; verify policy use before relying on it"),
        ("Solution Management", "ViewSolutionLibrary", "View solution library", "Seeded permission; verify policy use before relying on it"),
        ("Solution Management", "ImportExportCards", "Import/export cards", "Seeded permission; verify policy use before relying on it"),
        ("Analytics", "ViewAnalyticsDashboard", "View analytics dashboards", "Seeded permission; verify policy use before relying on it"),
    ]
    add_table(
        doc,
        ["Category", "Permission", "Plain-language capability", "Default / intended assignment"],
        permission_rows,
        widths=[1.3, 1.55, 2.1, 1.55],
        font_size=7.0,
    )

    add_heading(doc, "Annex 2 - Database Table Summary", 1)
    add_para(
        doc,
        "This annex gives the team a practical map of where information is stored. It is not "
        "a full entity relationship diagram, but it is enough for handover discussions, backup "
        "planning, data ownership reviews, and future technical orientation.",
    )
    add_bullets(
        doc,
        [
            "Use RosraReports as the first place to understand a report's profile, workflow status, and saved JSON payloads.",
            "Use AnalysisSnapshots and ReportArtifacts when checking what was preserved during submission or validation.",
            "Use AuditLogs, DataUploadHistory, and SolutionCardHistory when reconstructing operational changes.",
        ],
    )
    add_table(
        doc,
        ["Table / group", "Purpose"],
        [
            ("AspNetUsers and Identity tables", "Store user accounts, roles, claims, logins, tokens, and memberships."),
            ("Permissions and RolePermissions", "Store application-specific permission names and role-permission links."),
            ("RosraReports", "Central assessment report table with metadata, status, and JSON report payloads."),
            ("ReviewNotes", "Reviewer comments and revision notes attached to reports."),
            ("AnalysisSnapshots", "Stable serialized report copies created at submission, validation, or backup points."),
            ("ReportArtifacts", "Metadata for generated PDF and Excel files."),
            ("AuditLogs", "Operational and workflow audit trail."),
            ("Country, DB_Countries, DB_Frontiers", "Country, economic, government, and benchmark reference data."),
            ("Peers_SNG", "Admin-managed peer subnational government data."),
            ("UserPeerSngs", "User-owned peer SNG entries."),
            ("DataUploadHistory", "Record of uploaded reference data."),
            ("EmailSettings and EmailLogs", "SMTP configuration and email delivery records."),
            ("SystemSettings", "System-level key/value settings."),
            ("SolutionCards, SolutionCardHistory, CardSets", "Recommendation library, history, and grouped card sets."),
            ("__EFMigrationsHistory", "Entity Framework Core migration tracking."),
        ],
        widths=[2.2, 4.3],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Database ownership note",
        "The database backup protects SQL records, but it does not automatically protect "
        "generated files, platform secrets, Data Protection keys, or deployment configuration. "
        "Those items must be managed as part of the operating environment.",
        fill=PALE_ORANGE_HEX,
        accent=ORANGE,
    )

    add_heading(doc, "Annex 3 - Hosting and Configuration Checklist", 1)
    add_para(
        doc,
        "This checklist should be used during deployment, handover, and production support. "
        "It records which settings are required for the app to start, connect to the database, "
        "send email, support optional UN account sign-in, and generate PDFs reliably.",
    )
    add_para(
        doc,
        "The values themselves should not be written into this document. Store real secrets "
        "in Azure App Service settings, Azure DevOps secret variables, or an approved secret "
        "manager, and only record ownership and verification status here.",
    )
    add_table(
        doc,
        ["Configuration item", "Expected handling"],
        [
            ("CONNECTION_STRING", "Preferred production SQL Server/Azure SQL connection string."),
            ("ConnectionStrings:DefaultConnection", "Fallback .NET connection string."),
            ("AZURE_SQL_CONNECTIONSTRING", "Accepted Azure Service Connector/App Service connection string name."),
            ("ADMIN_SEED_PASSWORD", "Used only to create first admin account; rotate immediately after initial login."),
            ("EmailSettings:*", "SMTP host, port, sender, credentials, TLS, and notification toggles."),
            ("EntraId:TenantId", "Microsoft Entra tenant for optional SSO."),
            ("EntraId:ClientId", "Microsoft Entra app client ID."),
            ("EntraId:ClientSecret", "Microsoft Entra app secret stored in platform secrets."),
            ("EntraId:AdminEmails", "Comma-separated allow-list for admin role assignment from external login."),
            ("Data Protection keys", "Persist under HOME/DataProtection-Keys or equivalent persistent storage."),
            ("Playwright browser cache", "Persist/install under HOME/.playwright or equivalent location."),
            ("ASPNETCORE_ENVIRONMENT", "Production for hosted deployment."),
        ],
        widths=[2.1, 4.4],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Deployment verification note",
        "After configuration changes, verify the home page, login, Admin access, report "
        "creation, dashboard reload, review queue, PDF export, Excel export, and email test "
        "where SMTP is enabled.",
        fill=PALE_GREEN_HEX,
        accent=RGBColor(0x02, 0x7A, 0x48),
    )

    add_heading(doc, "Annex 4 - Screenshot Checklist", 1)
    add_para(
        doc,
        "Screenshots are included to make the handover document usable by non-technical team "
        "members. Before sharing the document externally, the team should replace any older "
        "or development screenshots with current production screenshots and confirm that no "
        "private user data is visible.",
    )
    add_table(
        doc,
        ["Document section", "Screenshot included / recommended"],
        [
            ("Part I", "Diagnostic workflow, user entry/login, system overview screens."),
            ("Part II", "Registration/login and admin navigation for user, role, and permission management."),
            ("Part III", "Quick estimate, detailed diagnostic, prioritisation, recommendations, and export modal."),
            ("Part IV", "Admin data management and expanded prioritisation/data views."),
            ("Part V", "Admin maintenance interface and operational management screens."),
            ("Future replacement", "Replace historical screenshots with current production screenshots before external publication."),
            ("Accessibility", "Keep captions under every screenshot and avoid screenshots as the only source of critical instructions."),
        ],
        widths=[1.8, 4.7],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Screenshot quality note",
        "Use clear browser screenshots that show the full workflow context. Avoid cropped "
        "screenshots that hide the page title, navigation, status, or action buttons needed "
        "to understand what the user is doing.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )

    add_heading(doc, "Annex 5 - Backup and Restore Checklist", 1)
    add_para(
        doc,
        "This annex is intended for operations and disaster recovery planning. The built-in "
        "Excel exports are useful for analysis, but they should not be treated as complete "
        "system backups. Full recovery requires database backup plus environment and file "
        "storage checks.",
    )
    add_para(
        doc,
        "The restore process should be tested against a non-production database before any "
        "production restore. After restore, the team should verify both data integrity and "
        "application behavior, especially login, report loading, review records, and exports.",
    )
    add_table(
        doc,
        ["Step", "Checklist item"],
        [
            ("1", "Confirm production database name and hosting type."),
            ("2", "Create SQL Server .bak or Azure SQL .bacpac backup before major deployments or migrations."),
            ("3", "Verify backup integrity or complete export status."),
            ("4", "Back up report artifact file storage separately from SQL."),
            ("5", "Back up Data Protection keys and record where they are stored."),
            ("6", "Confirm environment variables/secrets are recorded in approved secret management, not in the repository."),
            ("7", "Restore into a test database first when possible."),
            ("8", "Verify login, dashboard, report loading, review queue, admin pages, and exports after restore."),
            ("9", "Run dotnet ef database update only when application code is newer than restored schema."),
            ("10", "Document restore result and responsible person."),
        ],
        widths=[0.7, 5.8],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Restore control",
        "Do not overwrite a production database until a backup has been verified and a test "
        "restore has been inspected. Record the restore date, source backup, target database, "
        "migration status, and person responsible.",
        fill=PALE_RED_HEX,
        accent=RGBColor(0xB4, 0x23, 0x18),
    )

    add_heading(doc, "Annex 6 - Maintenance and Handover Checklist", 1)
    add_para(
        doc,
        "This checklist should be reviewed whenever the system is handed to a new team member, "
        "moved to a new hosting environment, or changed in a way that affects reports, users, "
        "permissions, data, exports, or configuration.",
    )
    add_para(
        doc,
        "For day-to-day operations, the most important discipline is to keep implementation, "
        "documentation, backups, and access control aligned. A system change is not complete "
        "until the relevant team procedures and handover notes have also been updated.",
    )
    add_table(
        doc,
        ["Area", "Team action"],
        [
            ("Production URL", "Record exact live URL and any custom domain."),
            ("Azure ownership", "Record subscription, resource group, App Service, database, and responsible owner."),
            ("Secrets", "Rotate any exposed token and keep all credentials in approved secret storage."),
            ("Admin account", "Confirm admin@rosra.com seed behavior and rotate the initial password after first login."),
            ("Roles", "Confirm Admin, Reviewer, and User role assignments match operating policy."),
            ("Permissions", "Review seeded permissions and decide whether solution-card permissions need active policy enforcement."),
            ("Database", "Confirm migrations are applied and backup schedule is active."),
            ("Artifacts", "Confirm where PDF/Excel artifacts are stored and backed up."),
            ("Email", "Send test email after SMTP configuration changes."),
            ("PDF export", "Confirm Playwright Chromium installation completes after deployments."),
            ("Localization", "Review English, French, and Spanish after changing user-facing text."),
            ("Documentation", "Update this handover report after major workflow, access, schema, hosting, or export changes."),
        ],
        widths=[1.7, 4.8],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Handover close-out note",
        "At the end of handover, the team should confirm four items in writing: who owns "
        "production administration, who owns the database backup, who owns future development "
        "changes, and where current credentials and environment settings are stored.",
        fill=LIGHT_BLUE_HEX,
        accent=BLUE,
    )


def extract_reference_images(tmp_dir: Path):
    if REFERENCE_DOCX.exists():
        with zipfile.ZipFile(REFERENCE_DOCX) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    zf.extract(name, tmp_dir)
    media = tmp_dir / "word" / "media"
    return {
        "workflow": media / "image3.png",
        "register": media / "image15.png",
        "login": media / "image16.png",
        "quick_profile": media / "image17.png",
        "quick_results": media / "image19.png",
        "gap_analysis": media / "image20.png",
        "prioritisation": media / "image9.png",
        "recommendations": media / "image22.png",
        "export_modal": media / "image12.png",
        "admin_data": ROOT / "wwwroot/images/admin-guide/data-management-streams.png",
        "admin_priority": ROOT / "wwwroot/images/admin-guide/data-management-prioritization.png",
    }


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory(prefix="rosra-handover-assets-") as tmp:
        images = extract_reference_images(Path(tmp))
        doc = setup_document()
        cover_page(doc)
        front_matter(doc)
        part_i(doc, images)
        part_ii(doc, images)
        part_iii(doc, images)
        part_iv(doc, images)
        part_v(doc, images)
        annexes(doc)
        doc.core_properties.title = "ROSRA V2 Implementation, Workflow, Access, and Data Management Handover Report"
        doc.core_properties.subject = "ROSRA V2 implementation handover"
        doc.core_properties.author = PREPARED_BY
        doc.core_properties.keywords = "ROSRA, implementation, workflow, permissions, database, handover"
        doc.core_properties.created = datetime(2026, 7, 9, 10, 0, 0)
        doc.core_properties.modified = datetime(2026, 7, 9, 10, 0, 0)
        doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
